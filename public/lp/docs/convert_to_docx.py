#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
要件定義書をMarkdownからDOCX形式に変換するスクリプト
必要なライブラリ: python-docx, markdown2
インストール: pip install python-docx markdown2
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """ハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0000FF")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_requirement_document():
    """要件定義書をDOCX形式で作成"""
    
    # ドキュメントを作成
    doc = Document()
    
    # ドキュメントのプロパティ設定
    doc.core_properties.title = "INTERCONNECT システム要件定義書"
    doc.core_properties.author = "INTERCONNECT開発チーム"
    doc.core_properties.subject = "ビジネスマッチングプラットフォーム要件定義"
    
    # タイトルページ
    title = doc.add_heading('INTERCONNECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('システム要件定義書', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    version_info = doc.add_paragraph('Version 1.0')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_info = doc.add_paragraph('2025年1月31日')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 目次
    doc.add_heading('目次', 1)
    toc_items = [
        "1. プロジェクト概要",
        "2. システムアーキテクチャ",
        "3. 機能要件",
        "4. データベース設計",
        "5. UI/UXデザイン",
        "6. セキュリティ要件",
        "7. パフォーマンス要件",
        "8. 運用・保守",
        "9. 今後の拡張予定",
        "10. 開発スケジュール",
        "11. 成果物",
        "12. 制約事項",
        "13. 用語定義",
        "14. 改訂履歴"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # 1. プロジェクト概要
    doc.add_heading('1. プロジェクト概要', 1)
    
    doc.add_heading('1.1 システム名', 2)
    doc.add_paragraph('INTERCONNECT（インターコネクト）')
    
    doc.add_heading('1.2 システムの目的', 2)
    doc.add_paragraph(
        'ビジネスプロフェッショナル向けのマッチングプラットフォームを提供し、'
        '適切なビジネスパートナーや協業相手を見つけることを支援する。'
    )
    
    doc.add_heading('1.3 主要機能', 2)
    features = [
        'ユーザー認証・プロファイル管理',
        'AIを活用したマッチング機能',
        'リアルタイムメッセージング',
        'イベント管理',
        'ダッシュボード分析'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 2. システムアーキテクチャ
    doc.add_heading('2. システムアーキテクチャ', 1)
    
    doc.add_heading('2.1 技術スタック', 2)
    
    doc.add_heading('フロントエンド', 3)
    frontend_tech = [
        'HTML5, CSS3, JavaScript (ES6+)',
        'レスポンシブデザイン対応',
        'プログレッシブウェブアプリ（PWA）対応'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('バックエンド', 3)
    backend_tech = [
        'Supabase (PostgreSQL + リアルタイムサブスクリプション)',
        'Row Level Security (RLS) による認証・認可'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('インフラ', 3)
    infra_tech = [
        'GitHub (バージョン管理)',
        'Supabase Cloud (ホスティング)'
    ]
    for tech in infra_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    # 3. 機能要件
    doc.add_heading('3. 機能要件', 1)
    
    doc.add_heading('3.1 認証機能', 2)
    
    doc.add_heading('3.1.1 ユーザー登録', 3)
    doc.add_paragraph('機能概要: メールアドレスとパスワードによる新規登録')
    doc.add_paragraph('入力項目:')
    register_items = [
        'メールアドレス（必須）',
        'パスワード（必須、8文字以上）',
        '名前（必須）'
    ]
    for item in register_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('バリデーション:')
    validations = [
        'メールアドレス形式チェック',
        'パスワード強度チェック',
        '重複登録チェック'
    ]
    for validation in validations:
        doc.add_paragraph(validation, style='List Bullet')
    
    doc.add_heading('3.2 プロファイル機能', 2)
    
    doc.add_heading('3.2.1 プロファイル作成・編集', 3)
    doc.add_paragraph('基本情報:')
    profile_items = [
        '名前',
        'プロフィール画像',
        '役職（title）',
        '会社名（company）',
        '業界（industry）',
        '地域（location）',
        '自己紹介（bio）'
    ]
    for item in profile_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('スキル・興味:')
    skill_items = [
        'スキルタグ（複数選択可能）',
        '興味・関心タグ（複数選択可能）'
    ]
    for item in skill_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.3 マッチング機能', 2)
    
    doc.add_heading('3.3.1 マッチングアルゴリズム', 3)
    doc.add_paragraph('基本スコア計算:')
    
    # コードブロックの追加
    code_para = doc.add_paragraph()
    code_para.add_run('// スコア計算要素\n').font.name = 'Courier New'
    code_para.add_run('- プロファイル充実度: 15-30%\n').font.name = 'Courier New'
    code_para.add_run('- スキルマッチ: 最大15%\n').font.name = 'Courier New'
    code_para.add_run('- 地域一致: 8%\n').font.name = 'Courier New'
    code_para.add_run('- 業界一致: 8%\n').font.name = 'Courier New'
    code_para.add_run('- 興味の共通性: 最大6%\n').font.name = 'Courier New'
    code_para.add_run('- ランダム要素: ±10%（ガウス分布）').font.name = 'Courier New'
    
    doc.add_paragraph('スコア範囲: 15%〜95%')
    doc.add_paragraph('表示順序: スコアの高い順')
    
    doc.add_heading('3.3.2 マッチング表示', 3)
    doc.add_paragraph('カード形式表示:')
    card_items = [
        'プロフィール画像',
        '名前・役職・会社',
        'マッチングスコア（%表示）',
        'スキルタグ（最大3個表示）',
        'レーダーチャート（6軸評価）'
    ]
    for item in card_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('レーダーチャート項目:')
    radar_items = [
        'スキル充実度',
        '経験値',
        '地域適合性',
        '業界適合性',
        '活動レベル',
        '興味の一致度'
    ]
    for i, item in enumerate(radar_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')
    
    # 4. データベース設計
    doc.add_heading('4. データベース設計', 1)
    
    doc.add_heading('4.1 テーブル構造', 2)
    
    doc.add_heading('4.1.1 profiles テーブル', 3)
    
    # テーブル定義をコードブロックとして追加
    table_def = doc.add_paragraph()
    table_def.paragraph_format.left_indent = Inches(0.5)
    run = table_def.add_run(
        'CREATE TABLE profiles (\n'
        '    id UUID PRIMARY KEY REFERENCES auth.users(id),\n'
        '    name TEXT NOT NULL,\n'
        '    email TEXT NOT NULL,\n'
        '    avatar_url TEXT,\n'
        '    title TEXT,\n'
        '    company TEXT,\n'
        '    industry TEXT,\n'
        '    location TEXT,\n'
        '    bio TEXT,\n'
        '    skills TEXT[] DEFAULT \'{}\',\n'
        '    interests TEXT[] DEFAULT \'{}\',\n'
        '    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),\n'
        '    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()\n'
        ');'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # 5. UI/UXデザイン
    doc.add_heading('5. UI/UXデザイン', 1)
    
    doc.add_heading('5.1 デザイン原則', 2)
    design_principles = [
        'レスポンシブデザイン: モバイル、タブレット、デスクトップ対応',
        'アクセシビリティ: WCAG 2.1 AA準拠',
        'パフォーマンス: 初期表示3秒以内'
    ]
    for principle in design_principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    doc.add_heading('5.2 カラースキーム', 2)
    colors = [
        'プライマリカラー: #3498db（青）',
        'セカンダリカラー: #2ecc71（緑）',
        'アクセントカラー: #e74c3c（赤）',
        '背景色: #f8f9fa',
        'テキスト色: #2c3e50'
    ]
    for color in colors:
        doc.add_paragraph(color, style='List Bullet')
    
    # 6. セキュリティ要件
    doc.add_heading('6. セキュリティ要件', 1)
    
    doc.add_heading('6.1 認証・認可', 2)
    security_items = [
        'パスワード: bcryptによるハッシュ化',
        'セッション: JWT (JSON Web Token)',
        'HTTPS: 全通信をSSL/TLSで暗号化'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 7. パフォーマンス要件
    doc.add_heading('7. パフォーマンス要件', 1)
    
    doc.add_heading('7.1 レスポンスタイム', 2)
    performance_items = [
        'ページ読み込み: 3秒以内（3G回線）',
        'API応答: 500ms以内',
        '検索結果: 1秒以内'
    ]
    for item in performance_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 8. 運用・保守
    doc.add_heading('8. 運用・保守', 1)
    
    doc.add_heading('8.1 監視項目', 2)
    monitoring_items = [
        'システム稼働率: 99.9%以上',
        'エラー率: 0.1%以下',
        'レスポンスタイム: 継続監視'
    ]
    for item in monitoring_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 9. 今後の拡張予定
    doc.add_heading('9. 今後の拡張予定', 1)
    
    doc.add_heading('9.1 Phase 2（3ヶ月後）', 2)
    phase2_items = [
        'AIによる自動マッチング提案',
        'ビデオ通話機能',
        'カレンダー連携'
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('9.2 Phase 3（6ヶ月後）', 2)
    phase3_items = [
        'モバイルアプリ（iOS/Android）',
        '多言語対応（英語・中国語）',
        '有料プラン導入'
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 10. 開発スケジュール
    doc.add_heading('10. 開発スケジュール', 1)
    
    doc.add_heading('10.1 開発フェーズ', 2)
    schedule_items = [
        '要件定義: 2週間',
        '基本設計: 2週間',
        '詳細設計: 1週間',
        '実装: 6週間',
        'テスト: 2週間',
        'リリース準備: 1週間'
    ]
    for i, item in enumerate(schedule_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')
    
    doc.add_paragraph('合計: 約3ヶ月')
    
    # 11. 成果物
    doc.add_heading('11. 成果物', 1)
    
    doc.add_heading('11.1 ドキュメント', 2)
    documents = [
        '要件定義書（本書）',
        '基本設計書',
        '詳細設計書',
        'テスト仕様書',
        '運用マニュアル'
    ]
    for doc_item in documents:
        doc.add_paragraph(doc_item, style='List Bullet')
    
    doc.add_heading('11.2 ソースコード', 2)
    source_items = [
        'フロントエンド（HTML/CSS/JavaScript）',
        'データベーススキーマ（SQL）',
        '設定ファイル',
        'デプロイスクリプト'
    ]
    for item in source_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 12. 制約事項
    doc.add_heading('12. 制約事項', 1)
    
    doc.add_heading('12.1 技術的制約', 2)
    tech_constraints = [
        'ブラウザ対応：Chrome, Firefox, Safari, Edge（最新2バージョン）',
        'JavaScript有効必須',
        'Cookie有効必須'
    ]
    for constraint in tech_constraints:
        doc.add_paragraph(constraint, style='List Bullet')
    
    # 13. 用語定義
    doc.add_heading('13. 用語定義', 1)
    
    # テーブルを作成
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # ヘッダー行
    header_cells = table.rows[0].cells
    header_cells[0].text = '用語'
    header_cells[1].text = '説明'
    
    # データ行
    terms = [
        ('プロファイル', 'ユーザーの基本情報、スキル、興味などを含む情報'),
        ('マッチング', 'ユーザー間の適合度を算出し、推奨する機能'),
        ('コネクション', 'ユーザー間の繋がり、承認により成立'),
        ('スコア', 'マッチング度を示す百分率（15%〜95%）'),
        ('レーダーチャート', '6軸でユーザーの特性を視覚化したグラフ')
    ]
    
    for i, (term, description) in enumerate(terms, 1):
        cells = table.rows[i].cells
        cells[0].text = term
        cells[1].text = description
    
    # 14. 改訂履歴
    doc.add_heading('14. 改訂履歴', 1)
    
    # 改訂履歴テーブル
    revision_table = doc.add_table(rows=2, cols=4)
    revision_table.style = 'Table Grid'
    
    # ヘッダー行
    header_cells = revision_table.rows[0].cells
    header_cells[0].text = '版数'
    header_cells[1].text = '日付'
    header_cells[2].text = '改訂内容'
    header_cells[3].text = '作成者'
    
    # データ行
    data_cells = revision_table.rows[1].cells
    data_cells[0].text = '1.0'
    data_cells[1].text = '2025-01-31'
    data_cells[2].text = '初版作成'
    data_cells[3].text = 'INTERCONNECT開発チーム'
    
    # ドキュメントを保存
    output_path = os.path.join(os.path.dirname(__file__), 'INTERCONNECT_要件定義書.docx')
    doc.save(output_path)
    print(f"要件定義書を作成しました: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_requirement_document()